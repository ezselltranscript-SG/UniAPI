import os
import re
import logging
import zipfile
import rarfile
import patoolib
import tempfile
import subprocess
from typing import Dict, Any, List, Optional
from docxcompose.composer import Composer
from docx import Document

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ColumnMergerService:
    """Servicio para fusionar archivos Word manteniendo su formato original."""

    @staticmethod
    def extract_archive(archive_data: bytes, temp_dir: str) -> str:
        """Extrae un archivo comprimido en un directorio temporal.
        
        Args:
            archive_data: Datos binarios del archivo comprimido
            temp_dir: Directorio temporal donde extraer los archivos
            
        Returns:
            Ruta al directorio donde se extrajeron los archivos
        """
        # Guardar el archivo comprimido en el directorio temporal
        archive_path = os.path.join(temp_dir, "archive")
        with open(archive_path, "wb") as f:
            f.write(archive_data)
        
        # Intentar extraer como ZIP
        extract_dir = os.path.join(temp_dir, "extracted")
        os.makedirs(extract_dir, exist_ok=True)
        
        try:
            if zipfile.is_zipfile(archive_path):
                with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_dir)
                    logger.info(f"Archivo ZIP extraído en {extract_dir}")
            elif rarfile.is_rarfile(archive_path):
                with rarfile.RarFile(archive_path, 'r') as rar_ref:
                    rar_ref.extractall(extract_dir)
                    logger.info(f"Archivo RAR extraído en {extract_dir}")
            else:
                # Intentar con patoolib para otros formatos
                patoolib.extract_archive(archive_path, outdir=extract_dir)
                logger.info(f"Archivo extraído con patoolib en {extract_dir}")
        except Exception as e:
            logger.error(f"Error al extraer el archivo: {str(e)}")
            raise ValueError(f"Error al extraer el archivo: {str(e)}")
        
        return extract_dir

    @staticmethod
    def sort_files_by_part(files: List[str]) -> List[str]:
        """Ordena los archivos por número de parte en el nombre.
        
        Args:
            files: Lista de rutas de archivos
            
        Returns:
            Lista ordenada de rutas de archivos
        """
        def extract_part_number(filename):
            # Buscar patrones como 'part1', 'part01', 'part_1', etc.
            match = re.search(r'part[_-]?(\d+)', os.path.basename(filename).lower())
            if match:
                return int(match.group(1))
            return 0  # Si no hay número de parte, poner al principio
        
        return sorted(files, key=extract_part_number)

    @staticmethod
    def merge_word_documents(word_files: List[str], output_path: str) -> None:
        """Fusiona varios documentos Word en uno solo manteniendo el formato.
        
        Args:
            word_files: Lista de rutas a archivos Word
            output_path: Ruta donde guardar el documento fusionado
        """
        if not word_files:
            raise ValueError("No se encontraron archivos Word para fusionar")
        
        # Usar el primer documento como base
        master = Document(word_files[0])
        composer = Composer(master)
        
        # Añadir el resto de documentos con dos saltos de línea entre ellos
        for i, file_path in enumerate(word_files[1:], 1):
            try:
                # Añadir dos saltos de línea al final del documento actual
                # Esto crea una separación entre documentos
                last_paragraph = master.add_paragraph()
                last_paragraph.add_run().add_break()
                last_paragraph.add_run().add_break()
                
                # Añadir el siguiente documento
                doc = Document(file_path)
                composer.append(doc)
                logger.info(f"Documento añadido: {os.path.basename(file_path)}")
            except Exception as e:
                logger.error(f"Error al añadir documento {file_path}: {str(e)}")
                raise ValueError(f"Error al añadir documento {os.path.basename(file_path)}: {str(e)}")
        
        # Guardar el documento fusionado
        composer.save(output_path)
        logger.info(f"Documento fusionado guardado en {output_path}")

    @staticmethod
    def merge_documents(archive_data: bytes, output_filename: str, temp_dir: str) -> Dict[str, Any]:
        """Fusiona documentos Word de un archivo comprimido manteniendo el formato original.
        
        Args:
            archive_data: Datos binarios del archivo comprimido
            output_filename: Nombre del archivo resultante
            temp_dir: Directorio temporal para archivos intermedios
            
        Returns:
            Diccionario con información del archivo resultante
        """
        try:
            # Extraer el archivo comprimido
            extract_dir = ColumnMergerService.extract_archive(archive_data, temp_dir)
            
            # Buscar archivos Word
            word_files = []
            for root, _, files in os.walk(extract_dir):
                for file in files:
                    if file.lower().endswith(('.doc', '.docx')):
                        word_files.append(os.path.join(root, file))
            
            if not word_files:
                raise ValueError("No se encontraron archivos Word válidos.")
            
            # Ordenar archivos por número de parte
            word_files = ColumnMergerService.sort_files_by_part(word_files)
            
            # Fusionar los documentos Word
            if not output_filename.lower().endswith('.docx'):
                output_filename += '.docx'
                
            output_path = os.path.join(temp_dir, output_filename)
            ColumnMergerService.merge_word_documents(word_files, output_path)
            
            return {
                "output_file": output_path,
                "file_count": len(word_files),
                "success": True
            }
            
        except Exception as e:
            logger.error(f"Error al fusionar documentos: {str(e)}")
            raise ValueError(f"Error al fusionar documentos: {str(e)}")
    
    @staticmethod
    def convert_docx_to_pdf(docx_path: str, output_dir: str) -> Optional[str]:
        """
        Convierte un documento Word a PDF usando LibreOffice sin alterar el formato.
        
        Args:
            docx_path: Ruta al documento Word
            output_dir: Directorio para guardar el PDF
            
        Returns:
            Path to PDF file or None if error
        """
        try:
            # Asegurarse de que el directorio de salida existe
            os.makedirs(output_dir, exist_ok=True)
            
            # Construir el comando para LibreOffice
            # --headless: ejecutar sin interfaz gráfica
            # --convert-to pdf: convertir a PDF
            # --outdir: directorio de salida
            command = [
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', output_dir, docx_path
            ]
            
            logger.info(f"Ejecutando comando: {' '.join(command)}")
            
            # Ejecutar el comando
            process = subprocess.Popen(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            
            # Esperar a que termine el proceso
            stdout, stderr = process.communicate()
            
            # Verificar si hubo errores
            if process.returncode != 0:
                logger.error(f"Error al convertir documento: {stderr.decode()}")
                return None
            
            logger.info(f"Conversión exitosa: {stdout.decode()}")
            
            # Buscar el archivo PDF generado
            base_name = os.path.basename(docx_path)
            pdf_name = os.path.splitext(base_name)[0] + '.pdf'
            pdf_path = os.path.join(output_dir, pdf_name)
            
            if os.path.exists(pdf_path):
                logger.info(f"PDF generado: {pdf_path}")
                return pdf_path
            else:
                # Si no encontramos el archivo con el nombre esperado, buscar cualquier PDF
                files = os.listdir(output_dir)
                logger.info(f"Archivos en directorio: {files}")
                
                # Buscar cualquier PDF generado
                for file in files:
                    if file.endswith(".pdf"):
                        pdf_path = os.path.join(output_dir, file)
                        logger.info(f"PDF encontrado: {pdf_path}")
                        return pdf_path
                
                logger.error("No se encontró ningún PDF generado")
                return None
                
        except Exception as e:
            logger.error(f"Error en conversión: {str(e)}")
            return None
            
    @staticmethod
    def word_to_pdf(word_path: str) -> Dict[str, Any]:
        """
        Convierte un documento Word a PDF sin alterar el formato.
        
        Args:
            word_path: Ruta al documento Word
            
        Returns:
            Diccionario con información del archivo PDF resultante
        """
        try:
            # Crear un directorio temporal para la salida del PDF
            temp_dir = tempfile.mkdtemp()
            logger.info(f"Directorio temporal para PDF: {temp_dir}")
            
            # Obtener el nombre base del archivo
            base_name = os.path.splitext(os.path.basename(word_path))[0]
            
            # Convertir a PDF usando LibreOffice
            pdf_path = ColumnMergerService.convert_docx_to_pdf(word_path, temp_dir)
            
            if not pdf_path or not os.path.exists(pdf_path):
                logger.error(f"Error al convertir {word_path} a PDF")
                raise ValueError(f"Error al convertir el documento a PDF")
                
            logger.info(f"PDF generado correctamente: {pdf_path}")
            
            return {
                "pdf_path": pdf_path,
                "filename": f"{base_name}.pdf"
            }
            
        except Exception as e:
            logger.error(f"Error al convertir a PDF: {str(e)}")
            raise ValueError(f"Error al convertir a PDF: {str(e)}")
            
    @staticmethod
    def convert_word_bytes_to_pdf(word_data: bytes, filename: str, temp_dir: str) -> Dict[str, Any]:
        """
        Convierte un documento Word (en bytes) a PDF sin alterar el formato.
        
        Args:
            word_data: Datos binarios del documento Word
            filename: Nombre del archivo original
            temp_dir: Directorio temporal para archivos intermedios
            
        Returns:
            Diccionario con información del archivo PDF resultante
        """
        try:
            # Guardar el documento Word en el directorio temporal
            word_path = os.path.join(temp_dir, filename)
            with open(word_path, "wb") as f:
                f.write(word_data)
            
            logger.info(f"Documento Word guardado en: {word_path}")
            
            # Convertir el documento Word a PDF
            pdf_result = ColumnMergerService.word_to_pdf(word_path)
            
            return {
                "pdf_path": pdf_result["pdf_path"],
                "filename": pdf_result["filename"],
                "word_path": word_path,  # También devolvemos la ruta al Word por si es necesaria
                "success": True
            }
            
        except Exception as e:
            logger.error(f"Error al convertir Word a PDF: {str(e)}")
            raise ValueError(f"Error al convertir Word a PDF: {str(e)}")
    
    @staticmethod
    def merge_and_convert_to_pdf(archive_data: bytes, output_filename: str, temp_dir: str) -> Dict[str, Any]:
        """
        Fusiona documentos Word de un archivo comprimido y los convierte a PDF sin alterar el formato.
        
        Args:
            archive_data: Datos binarios del archivo comprimido
            output_filename: Nombre del archivo resultante
            temp_dir: Directorio temporal para archivos intermedios
            
        Returns:
            Diccionario con información del archivo PDF resultante
        """
        try:
            # Primero fusionar los documentos Word
            result = ColumnMergerService.merge_documents(archive_data, output_filename, temp_dir)
            
            # Obtener la ruta del documento Word fusionado
            word_path = result["output_file"]
            
            # Convertir el documento Word fusionado a PDF
            pdf_result = ColumnMergerService.word_to_pdf(word_path)
            
            return {
                "pdf_path": pdf_result["pdf_path"],
                "filename": pdf_result["filename"],
                "word_path": word_path,  # También devolvemos la ruta al Word por si es necesaria
                "success": True
            }
            
        except Exception as e:
            logger.error(f"Error al fusionar y convertir a PDF: {str(e)}")
            raise ValueError(f"Error al fusionar y convertir a PDF: {str(e)}")
