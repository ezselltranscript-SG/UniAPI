import os
import tempfile
import shutil
import logging
from typing import Dict, Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocxGeneratorService:
    """Service class to generate a Word document (.docx) from input variables
    with a predefined template and style similar to the provided JavaScript example.
    """

    @staticmethod
    def generate_docx(data: Dict[str, Any]) -> Dict[str, Any]:
        """Create a .docx file given variables.

        Expected keys in data:
        - city (str)
        - authorName (str)
        - date (str)
        - body (str)
        - documentName (str) optional, final name for the file (with or without .docx)
        """
        temp_dir = tempfile.mkdtemp()
        logger.info(f"Directorio temporal creado: {temp_dir}")

        try:
            # Defaults and derived values
            city = (data.get("city") or "City").strip()
            author = (data.get("authorName") or "Name").strip()
            date_str = (data.get("date") or "No Date").strip()
            body = (data.get("body") or "").strip()
            title = f"{city} – {author}"
            date_body = f"{date_str} – {body}" if body else date_str

            # File name handling
            fname = (data.get("documentName") or "document").strip() or "document"
            if not fname.lower().endswith(".docx"):
                fname = f"{fname}.docx"

            output_path = os.path.join(temp_dir, fname)

            # Build the document
            doc = Document()

            # Page setup: Letter portrait (8.5 x 11 inches) with 1" margins
            section = doc.sections[0]
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Title paragraph (centered, bold, Times New Roman, 12pt)
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run(title)
            run_title.bold = True
            run_title.font.name = "Times New Roman"
            run_title.font.size = Pt(12)  # 24 half-points in JS example
            # Remove any extra spacing after the title so body follows immediately
            p_title_format = p_title.paragraph_format
            p_title_format.space_after = Pt(0)
            p_title_format.space_before = Pt(0)

            # Body paragraph (justified, first-line indent 0.2", Times New Roman, 12pt)
            p_body = doc.add_paragraph()
            p_body_format = p_body.paragraph_format
            p_body_format.first_line_indent = Inches(0.2)
            # Ensure no extra space before body so there is no gap after title
            p_body_format.space_before = Pt(0)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_body = p_body.add_run(date_body)
            run_body.font.name = "Times New Roman"
            run_body.font.size = Pt(12)

            # Save
            doc.save(output_path)
            logger.info(f"Documento DOCX generado en: {output_path}")

            return {
                "docx_path": output_path,
                "filename": fname,
                "temp_dir": temp_dir,
            }
        except Exception as e:
            shutil.rmtree(temp_dir, ignore_errors=True)
            logger.error(f"Error al generar DOCX: {e}")
            raise
