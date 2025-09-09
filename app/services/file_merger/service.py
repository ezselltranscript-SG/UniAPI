import os
import re
import shutil
import tempfile
import zipfile
import rarfile
import patoolib
import logging
from typing import List, Dict, Any, Optional
from pypdf import PdfWriter, PdfReader
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION
from copy import deepcopy

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileMergerService:
    """Service class for merging PDF and DOCX files"""
    
    @staticmethod
    def extract_part_number(filename: str) -> int:
        """Extract part number from filename for sorting"""
        match = re.search(r'part(\d+)', filename.lower())
        return int(match.group(1)) if match else 0

    @staticmethod
    def sort_files_by_part(files: List[str]) -> List[str]:
        """Sort files by part number in filename"""
        return sorted(files, key=lambda f: FileMergerService.extract_part_number(os.path.basename(f)))

    @staticmethod
    def extract_compressed_file(file_path: str, extract_dir: str) -> List[str]:
        """Extract files from compressed archive (zip, rar, etc.)"""
        extracted_files = []
        
        # Try ZIP format
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
                return [os.path.join(extract_dir, name) for name in zip_ref.namelist() if not name.endswith('/')]
        except zipfile.BadZipFile:
            pass

        # Try RAR format
        try:
            rarfile.UNRAR_TOOL = 'unrar'
            with rarfile.RarFile(file_path) as rar_ref:
                rar_ref.extractall(extract_dir)
                return [os.path.join(extract_dir, name) for name in rar_ref.namelist() if not rar_ref.getinfo(name).isdir()]
        except rarfile.NotRarFile:
            pass

        # Try other formats using patool
        try:
            patoolib.extract_archive(file_path, outdir=extract_dir)
            for root, _, files in os.walk(extract_dir):
                for file in files:
                    extracted_files.append(os.path.join(root, file))
            return extracted_files
        except Exception as e:
            logger.error(f"Error extracting archive: {str(e)}")
            raise Exception(f"Error extracting archive: {str(e)}")

    @staticmethod
    def filter_files_by_extension(file_paths: List[str], extensions: List[str]) -> List[str]:
        """Filter files by extension"""
        return [p for p in file_paths if os.path.splitext(p)[1].lower() in extensions]

    @staticmethod
    def merge_pdf_files(file_paths: List[str], output_path: str) -> None:
        """Merge multiple PDF files into one"""
        merger = PdfWriter()
        for path in file_paths:
            reader = PdfReader(path)
            for page in reader.pages:
                merger.add_page(page)
        with open(output_path, "wb") as out:
            merger.write(out)

    @staticmethod
    def merge_docx_simple(file_paths: List[str], output_path: str) -> None:
        """
        Merge multiple DOCX files into one.
        Each document is inserted as a new section with its own header
        and starts on a new page.
        """
        if not file_paths:
            raise Exception("No DOCX files provided")
        
        if len(file_paths) == 1:
            # If there's only one file, just copy it
            shutil.copy(file_paths[0], output_path)
            return
        
        try:
            # Helper functions to copy header/footer content
            def _clear_paragraphs(container):
                for p in list(container.paragraphs):
                    p._element.getparent().remove(p._element)

            def _copy_paragraphs(src_container, dst_container):
                for sp in src_container.paragraphs:
                    dp = dst_container.add_paragraph()
                    dp.style = sp.style
                    for run in sp.runs:
                        dr = dp.add_run(run.text)
                        dr.bold = run.bold
                        dr.italic = run.italic
                        dr.underline = run.underline
                        dr.font.name = run.font.name
                        dr.font.size = run.font.size
                if not src_container.paragraphs:
                    dst_container.add_paragraph("")

            # Create a base document with the first file
            master = Document(file_paths[0])

            # Ensure first section does not try to link headers forward/back
            for sec in master.sections:
                try:
                    sec.header.is_linked_to_previous = False
                    sec.footer.is_linked_to_previous = False
                except Exception:
                    pass

            # Append each additional document; avoid manual page breaks
            for i, file_path in enumerate(file_paths[1:], 1):
                logger.info(f"Adding document {i}: {os.path.basename(file_path)}")
                src_doc = Document(file_path)

                # Insert a definite Section Break (Next Page) so the next
                # letter starts on a new page
                master.add_section(WD_SECTION.NEW_PAGE)
                target_section = master.sections[-1]

                # Configure the new section's header/footer by copying from source
                target_section.header.is_linked_to_previous = False
                target_section.footer.is_linked_to_previous = False

                try:
                    _clear_paragraphs(target_section.header)
                    _copy_paragraphs(src_doc.sections[0].header, target_section.header)
                except Exception as he:
                    logger.warning(f"Could not copy header from {os.path.basename(file_path)}: {he}")

                try:
                    _clear_paragraphs(target_section.footer)
                    _copy_paragraphs(src_doc.sections[0].footer, target_section.footer)
                except Exception as fe:
                    logger.warning(f"Could not copy footer from {os.path.basename(file_path)}: {fe}")

                # Append the source document's body elements directly, but insert
                # them BEFORE the sectPr of the new section so the content belongs
                # to this section (and uses its header/footer)
                src_body = src_doc.element.body
                dst_body = master.element.body
                # Find index of sectPr (usually last). If not found, append at end.
                dst_children = list(dst_body)
                sectpr_index = None
                for idx, child in enumerate(dst_children[::-1]):
                    if child.tag.endswith('sectPr'):
                        sectpr_index = len(dst_children) - 1 - idx
                        break
                insert_index = sectpr_index if sectpr_index is not None else len(dst_children)
                for child in list(src_body):
                    if child.tag.endswith('sectPr'):
                        continue
                    dst_body.insert(insert_index, deepcopy(child))
                    insert_index += 1

            # Save the combined document
            master.save(output_path)
            logger.info(f"Combined document saved to {output_path}")
            
        except Exception as e:
            logger.error(f"Error merging documents: {str(e)}")
            raise Exception(f"Error merging documents: {str(e)}")
    
    @staticmethod
    def merge_files(archive_data: bytes, output_filename: str, temp_dir: str) -> Dict[str, Any]:
        """
        Main method to merge files from an archive
        
        Args:
            archive_data: Bytes of the archive file
            output_filename: Name for the output file
            temp_dir: Temporary directory for processing
            
        Returns:
            Dictionary with output path and media type
        """
        # Save the uploaded archive to a temporary file
        temp_path = os.path.join(temp_dir, "archive_input")
        with open(temp_path, "wb") as f:
            f.write(archive_data)

        # Extract the archive
        extract_dir = os.path.join(temp_dir, "extracted")
        os.makedirs(extract_dir, exist_ok=True)
        extracted = FileMergerService.extract_compressed_file(temp_path, extract_dir)

        # Filter for PDF and DOCX files
        valid_exts = ['.pdf', '.docx']
        filtered = FileMergerService.filter_files_by_extension(extracted, valid_exts)
        if not filtered:
            raise Exception("No valid PDF or DOCX files found.")

        # Sort files by part number
        sorted_files = FileMergerService.sort_files_by_part(filtered)
        
        # Determine the output file type based on the first file
        ext = os.path.splitext(sorted_files[0])[1].lower()
        output_path = os.path.join(temp_dir, f"{output_filename}{ext}")

        # Merge files based on type
        if ext == ".pdf":
            FileMergerService.merge_pdf_files(sorted_files, output_path)
            media_type = "application/pdf"
        elif ext == ".docx":
            FileMergerService.merge_docx_simple(sorted_files, output_path)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            raise Exception("Unsupported file type.")

        return {
            "output_path": output_path,
            "filename": f"{output_filename}{ext}",
            "media_type": media_type
        }
