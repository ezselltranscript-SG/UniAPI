import io
import json
import logging
import os
import re
import time
import urllib.parse
import urllib.request
from typing import Any, Dict, List, Tuple

logger = logging.getLogger(__name__)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


_NORM_CACHE: Dict[str, Any] = {"by_client": {}}
_NORM_TTL_SECONDS = 300
_CONTEXT_WINDOW = 80


def _fetch_normalization_rules(client: str) -> List[Tuple[str, str, str]]:
    """Returns list of (short_form, expansion, notes)."""
    supabase_url = (os.getenv("SUPABASE_URL") or "").strip().rstrip("/")
    supabase_key = (os.getenv("SUPABASE_KEY") or "").strip()
    if not supabase_url or not supabase_key:
        return []

    now = time.time()
    by_client = _NORM_CACHE.get("by_client") or {}
    entry = by_client.get(client) or {}
    if now - float(entry.get("ts") or 0.0) < _NORM_TTL_SECONDS:
        return list(entry.get("items") or [])

    headers = {
        "apikey": supabase_key,
        "Authorization": f"Bearer {supabase_key}",
        "Accept": "application/json",
    }

    rows = None
    for columns in ("ShortForm,Expansion,Notes,Client", "ShortForm,Expansion,Client"):
        query = urllib.parse.urlencode({"select": columns})
        url = f"{supabase_url}/rest/v1/Normalization_Words_List?{query}"
        req = urllib.request.Request(url, headers=headers, method="GET")
        try:
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            if isinstance(data, list):
                rows = data
                break
        except Exception:
            continue

    if rows is None:
        return []

    items: List[Tuple[str, str, str]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        row_client = (row.get("Client") or "").strip()
        if client.lower() not in row_client.lower():
            continue
        short_form = (row.get("ShortForm") or "").strip()
        expansion = (row.get("Expansion") or "").strip()
        notes = (row.get("Notes") or "").strip()
        if short_form and expansion:
            items.append((short_form, expansion, notes))

    if not isinstance(_NORM_CACHE.get("by_client"), dict):
        _NORM_CACHE["by_client"] = {}
    _NORM_CACHE["by_client"][client] = {"ts": now, "items": items}
    return items


def _extract_alternate_forms(notes: str, short_form: str) -> List[str]:
    """Pull dotted/alternate forms out of the notes (e.g. 'A.M.', 'P.M.')."""
    if not notes:
        return []
    alts = re.findall(r'[A-Za-z](?:\.[A-Za-z])+\.?', notes)
    return [a for a in alts if a.lower() != short_form.lower()]


def _expand_rules_with_alternates(rules: List[Tuple[str, str, str]]) -> List[Tuple[str, str, str]]:
    """For each rule, append extra entries for any alternate forms found in its notes."""
    expanded: List[Tuple[str, str, str]] = []
    for short_form, expansion, notes in rules:
        expanded.append((short_form, expansion, notes))
        for alt in _extract_alternate_forms(notes, short_form):
            expanded.append((alt, expansion, notes))
    return expanded


def _build_pattern(short_form: str) -> re.Pattern:
    escaped = re.escape(short_form)
    if re.fullmatch(r"[A-Za-z]+", short_form):
        # Allow matching after digits (e.g. "73deg" → "73 degrees").
        # (?<![A-Za-z]) prevents matching inside other words while still
        # catching the word immediately after a number.
        return re.compile(rf"(?<![A-Za-z]){escaped}(?![A-Za-z0-9])", re.IGNORECASE)
    if re.fullmatch(r"[A-Za-z0-9']+", short_form):
        return re.compile(rf"\b{escaped}\b", re.IGNORECASE)
    return re.compile(escaped, re.IGNORECASE)


def _context_is_uppercase(text: str) -> bool:
    alpha = [c for c in text if c.isalpha()]
    if not alpha:
        return False
    return sum(1 for c in alpha if c.isupper()) / len(alpha) > 0.5


def _get_context(text: str, start: int, end: int) -> str:
    before = text[max(0, start - _CONTEXT_WINDOW):start]
    matched = text[start:end]
    after = text[end:end + _CONTEXT_WINDOW]
    return f"{before}[{matched}]{after}"


def _build_prompt(text: str, matches: List[Dict[str, Any]]) -> str:
    match_lines = []
    for m in matches:
        notes_text = (
            f"\n  NOTES (these take priority — follow them precisely): {m['notes']}"
            if m["notes"] else ""
        )
        match_lines.append(
            f"ID {m['id']}: \"{m['short_form']}\" → \"{m['expansion']}\""
            f"{notes_text}"
            f"\n  Context snippet: {m['context']}"
        )
    return (
        "You are processing a correspondence letter. For each potential word expansion below, "
        "decide whether the short form should be replaced with its expansion based on its context.\n\n"
        "Decision rules (apply in order — stop at the first rule that matches):\n"
        "1. If NOTES are provided for a match, follow them precisely — they override all rules below.\n"
        "2. Do NOT expand when the word is a verb, even when the subject is omitted "
        "(e.g. 'I am', 'but am sure', 'am going', 'am not').\n"
        "3. Do NOT expand AM/PM or A.M./P.M. when directly preceded by a clock time "
        "(e.g. '7:30 AM', '9:00 PM' — leave these as-is).\n"
        "4. DO expand AM, PM, A.M., P.M. when they stand alone or follow a day/period word "
        "(e.g. 'Sunday PM', 'Monday AM', 'a.m.' by itself).\n"
        "5. Do NOT expand when the word refers to something other than the intended expansion "
        "(e.g. 'sun' as the celestial body should not expand to 'Sunday').\n"
        "6. Do NOT expand proper names.\n"
        "7. Only expand when you are confident the word is being used as an abbreviation. "
        "If genuinely uncertain, do NOT expand.\n\n"
        f"Full letter text:\n\"\"\"\n{text}\n\"\"\"\n\n"
        "Potential expansions:\n" + "\n\n".join(match_lines) + "\n\n"
        "Return ONLY valid JSON:\n"
        '{"decisions": [{"id": <int>, "expand": <bool>}]}'
    )


def _parse_ai_response(response_text: str, matches: List[Dict[str, Any]]) -> Dict[int, bool]:
    json_match = re.search(r"\{.*\}", response_text, re.DOTALL)
    if not json_match:
        return {m["id"]: True for m in matches}
    data = json.loads(json_match.group())
    decisions = {int(d["id"]): bool(d["expand"]) for d in data.get("decisions", [])}
    skipped = [m["short_form"] for m in matches if not decisions.get(m["id"], True)]
    expanded = [m["short_form"] for m in matches if decisions.get(m["id"], True)]
    logger.info("AI decisions — expand: %s | skip: %s", expanded, skipped)
    return decisions


def _ai_context_decisions(text: str, matches: List[Dict[str, Any]]) -> Dict[int, bool]:
    """
    Ask an LLM whether each match should be expanded based on context.

    Provider selection (checked in order):
      1. Anthropic — if ANTHROPIC_API_KEY is set in the environment.
      2. OpenAI    — if OPENAI_API_KEY is set in the environment.
      3. No AI     — all matches default to expand=True.

    Returns {match_id: should_expand}.
    """
    if not matches:
        return {}

    anthropic_key = (os.getenv("ANTHROPIC_API_KEY") or "").strip()
    openai_key = (os.getenv("OPENAI_API_KEY") or "").strip()

    if not anthropic_key and not openai_key:
        logger.warning("No AI API key found (ANTHROPIC_API_KEY / OPENAI_API_KEY) — all matches will expand")
        return {m["id"]: True for m in matches}

    prompt = _build_prompt(text, matches)

    # --- Anthropic (preferred) ---
    if anthropic_key:
        try:
            import anthropic
            logger.info("AI context check: sending %d match(es) to Anthropic claude-haiku", len(matches))
            ai_client = anthropic.Anthropic(api_key=anthropic_key)
            message = ai_client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}],
            )
            return _parse_ai_response(message.content[0].text.strip(), matches)
        except ImportError:
            logger.warning("anthropic package not installed — trying OpenAI fallback")
        except Exception as exc:
            logger.error("Anthropic AI check failed (%s) — trying OpenAI fallback", exc)

    # --- OpenAI (fallback) ---
    if openai_key:
        try:
            import openai as openai_module
            logger.info("AI context check: sending %d match(es) to OpenAI gpt-4o-mini", len(matches))
            oa_client = openai_module.OpenAI(api_key=openai_key)
            completion = oa_client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}],
            )
            return _parse_ai_response(completion.choices[0].message.content.strip(), matches)
        except ImportError:
            logger.warning("openai package not installed — skipping AI context check")
        except Exception as exc:
            logger.error("OpenAI AI check failed (%s) — falling back to expand all", exc)

    return {m["id"]: True for m in matches}


def _apply_rules(text: str, rules: List[Tuple[str, str, str]]) -> Tuple[str, List[Dict[str, Any]]]:
    """Apply normalization rules with AI context checking. Returns (normalized_text, changes_list)."""
    uppercase_context = _context_is_uppercase(text)

    # Collect all potential matches from the original text for a single batched AI call.
    # Keyed by (short_form_lower, occurrence_index) so decisions survive sequential substitutions.
    all_match_requests: List[Dict[str, Any]] = []
    occurrence_to_id: Dict[Tuple[str, int], int] = {}
    match_id = 0

    for short_form, expansion, notes in rules:
        pattern = _build_pattern(short_form)
        for occ_idx, m in enumerate(pattern.finditer(text)):
            all_match_requests.append({
                "id": match_id,
                "short_form": short_form,
                "expansion": expansion,
                "notes": notes,
                "context": _get_context(text, m.start(), m.end()),
            })
            occurrence_to_id[(short_form.lower(), occ_idx)] = match_id
            match_id += 1

    # Single AI call for all matches in this text block
    all_decisions = _ai_context_decisions(text, all_match_requests)

    expand_map: Dict[Tuple[str, int], bool] = {
        key: all_decisions.get(mid, True)
        for key, mid in occurrence_to_id.items()
    }

    # Apply rules sequentially using the saved decisions.
    # Occurrence indices stay stable because well-formed abbreviation rules don't
    # introduce new instances of each other's short forms.
    changes: List[Dict[str, Any]] = []
    normalized = text

    for short_form, expansion, notes in rules:
        pattern = _build_pattern(short_form)
        applied = expansion.upper() if uppercase_context else expansion.lower()
        sf_lower = short_form.lower()

        current_matches = list(pattern.finditer(normalized))
        if not current_matches:
            continue

        parts: List[str] = []
        last = 0
        count = 0

        for occ_idx, m in enumerate(current_matches):
            if expand_map.get((sf_lower, occ_idx), True):
                parts.append(normalized[last:m.start()])
                # Insert a space when the match is directly adjacent to alphanumeric text
                # (e.g. "73deg" → "73 degrees" instead of "73degrees").
                needs_space = m.start() > 0 and normalized[m.start() - 1].isalnum()
                parts.append((" " if needs_space else "") + applied)
                count += 1
            else:
                parts.append(normalized[last:m.end()])
            last = m.end()

        parts.append(normalized[last:])
        normalized = "".join(parts)

        if count > 0:
            changes.append({"short_form": short_form, "expansion": expansion, "occurrences": count})

    return normalized, changes


def normalize_text(date: str, body: str, client: str) -> Dict[str, Any]:
    """
    Apply normalization rules to both date and body.

    Returns:
      - normalized_date: str
      - normalized_body: str
      - changes: list of {short_form, expansion, occurrences}
      - total_changes: int
    """
    rules = _expand_rules_with_alternates(_fetch_normalization_rules(client))

    normalized_date, date_changes = _apply_rules(date, rules)
    normalized_body, body_changes = _apply_rules(body, rules)

    merged: Dict[str, Dict[str, Any]] = {}
    for change in date_changes + body_changes:
        key = change["short_form"]
        if key in merged:
            merged[key]["occurrences"] += change["occurrences"]
        else:
            merged[key] = dict(change)
    changes = list(merged.values())

    total_changes = sum(c["occurrences"] for c in changes)
    return {
        "normalized_date": normalized_date,
        "normalized_body": normalized_body,
        "changes": changes,
        "total_changes": total_changes,
    }


def build_report_docx(changes: List[Dict[str, Any]], total_changes: int) -> bytes:
    """Generate a Word document listing all word substitutions and the total count."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Normalization Report")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    header = doc.add_paragraph()
    h_run = header.add_run("Words Changed:")
    h_run.bold = True
    h_run.font.size = Pt(12)
    h_run.font.name = "Times New Roman"

    if changes:
        for change in changes:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(
                f'"{change["short_form"]}"  →  "{change["expansion"]}"'
                f'  ({change["occurrences"]} occurrence(s))'
            )
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"
    else:
        p = doc.add_paragraph()
        r = p.add_run("No changes made.")
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"

    doc.add_paragraph()

    total_p = doc.add_paragraph()
    t_run = total_p.add_run(f"Total words changed: {total_changes}")
    t_run.bold = True
    t_run.font.size = Pt(12)
    t_run.font.name = "Times New Roman"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
