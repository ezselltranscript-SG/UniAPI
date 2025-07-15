import os
import tempfile
import shutil
import subprocess
import logging
import io
import re
from typing import Dict, Any, Tuple, Optional
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfWriter, PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class WordToPdfService:
    """Service class for converting Word documents to PDF with advanced header formatting"""
    
    @staticmethod
    def modify_document_headers(word_path: str) -> Tuple[str, Optional[str]]:
        """
        Modifica los encabezados del documento Word para que cada página tenga el formato correcto
        
        Args:
            word_path: Ruta al documento Word
            
        Returns:
            Tuple with path to modified document and base code
        """
        try:
            # Extraer el nombre base del archivo
            original_filename = os.path.basename(word_path)
            temp_dir = tempfile.mkdtemp()
            base_name = os.path.basename(word_path)
            modified_docx = os.path.join(temp_dir, f"modified_{base_name}")
            
            # Abrir el documento original
            doc = Document(word_path)
            
            # Extraer el código base del nombre del archivo exactamente como aparece
            # Ejemplo: "062725-0620-b04-25.docx" -> "062725-0620-b04-25"
            base_code = os.path.splitext(base_name)[0]
            logger.info(f"Código base identificado: {base_code}")
            
            # Si no se encuentra un código base, usar un valor predeterminado
            if not base_code:
                base_code = "transcript"
                logger.warning(f"No se identificó código base, usando valor predeterminado: {base_code}")
            
            # ELIMINAR COMPLETAMENTE los encabezados de cada sección
            for section_idx, section in enumerate(doc.sections):
                part_number = section_idx + 1
                header = section.header
                
                # Eliminar todo el contenido del encabezado
                for paragraph in list(header.paragraphs):
                    p = paragraph._element
                    p.getparent().remove(p)
                    paragraph._p = None
                    paragraph._element = None
                
                # Añadir un párrafo vacío para mantener la estructura
                header.add_paragraph()
                
                logger.info(f"Eliminado encabezado para sección {part_number}")
            
            # Forzar Times New Roman 10 en todos los estilos
            try:
                from docx.enum.style import WD_STYLE_TYPE
                for style in doc.styles:
                    if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
                        if style.font is not None:
                            style.font.name = 'Times New Roman'
                            style.font.size = Pt(10)
            except Exception as e:
                logger.warning(f"No se pudo modificar estilos globales: {e}")

            # Cambiar la fuente y tamaño manualmente en cada ejecución de texto
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            
            # Cambiar también en las tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10)
            
            # Guardar el documento modificado
            doc.save(modified_docx)
            logger.info(f"Documento con encabezados eliminados guardado en: {modified_docx}")
            
            return modified_docx, base_code
            
        except Exception as e:
            logger.error(f"Error al modificar encabezados: {str(e)}")
            return None, None
    
    @staticmethod
    def add_page_headers_to_pdf(pdf_path: str, base_code: str) -> Optional[str]:
        """
        Modifica un PDF para añadir encabezados diferentes a cada página
        con el formato exacto base_code_Part1, base_code_Part2, etc.
        Mantiene el código base exactamente como está, sin modificarlo.
        
        Args:
            pdf_path: Ruta al archivo PDF
            base_code: Código base para los encabezados
            
        Returns:
            Path to modified PDF or None if error
        """
        try:
            # Crear un nuevo PDF con los encabezados correctos
            output_pdf = os.path.join(os.path.dirname(pdf_path), f"headers_{os.path.basename(pdf_path)}")
            
            # Abrir el PDF original
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            # Para cada página, añadir el encabezado correcto
            for i, page in enumerate(reader.pages):
                # Crear un PDF en memoria con el encabezado
                packet = io.BytesIO()
                can = canvas.Canvas(packet, pagesize=letter)
                
                # Dibujar un rectángulo blanco para cubrir completamente cualquier encabezado existente
                can.setFillColorRGB(1, 1, 1)  # Color blanco
                # Bajar el rectángulo y el texto unos 20 puntos
                can.rect(0, 750, 612, 28, fill=True, stroke=False)  # 750 en vez de 770
                
                # Configurar el encabezado con el número de parte correcto
                part_number = i + 1
                header_text = f"{base_code}_Part{part_number}"
                
                # Añadir el texto del encabezado en la posición correcta (esquina superior izquierda, pero más abajo)
                can.setFillColorRGB(0, 0, 0)  # Color negro para el texto
                can.setFont("Helvetica", 10)
                can.drawString(25, 765, header_text)  # 765 en vez de 785
                can.save()
                
                # Mover al inicio del BytesIO
                packet.seek(0)
                watermark = PdfReader(packet)
                
                # Fusionar la página original con el encabezado
                page.merge_page(watermark.pages[0])
                writer.add_page(page)
                
                logger.info(f"Añadido encabezado a página {part_number}: {header_text}")
            
            # Guardar el PDF modificado
            with open(output_pdf, "wb") as output_stream:
                writer.write(output_stream)
            
            logger.info(f"PDF con encabezados modificados guardado en: {output_pdf}")
            
            # Reemplazar el PDF original con el modificado
            shutil.move(output_pdf, pdf_path)
            
            return pdf_path
        
        except Exception as e:
            logger.error(f"Error al añadir encabezados al PDF: {str(e)}")
            return None
    
    @staticmethod
    def convert_docx_to_pdf(docx_path: str, output_dir: str) -> Optional[str]:
        """
        Convierte un documento Word a PDF usando LibreOffice de manera simple.
        
        Args:
            docx_path: Ruta al documento Word
            output_dir: Directorio para guardar el PDF
            
        Returns:
            Path to PDF file or None if error
        """
        try:
            # Asegurarse de que output_dir sea un directorio válido
            os.makedirs(output_dir, exist_ok=True)
            
            # Nombre base del archivo sin extensión
            base_name = os.path.basename(docx_path)
            base_name_without_ext = os.path.splitext(base_name)[0]
            expected_pdf = os.path.join(output_dir, f"{base_name_without_ext}.pdf")
            
            # Comando simple para convertir a PDF
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf:writer_pdf_Export",
                "--outdir", output_dir,
                docx_path
            ]
            
            logger.info(f"Ejecutando: {' '.join(cmd)}")
            
            # Ejecutar el comando
            process = subprocess.run(cmd, capture_output=True, text=True)
            
            # Registrar la salida
            if process.stdout:
                logger.info(f"Salida: {process.stdout}")
            if process.stderr:
                logger.warning(f"Error: {process.stderr}")
            
            # Verificar el archivo PDF generado
            if os.path.exists(expected_pdf):
                logger.info(f"PDF generado correctamente: {expected_pdf}")
                return expected_pdf
            else:
                # Listar archivos en el directorio para diagnóstico
                files = os.listdir(output_dir)
                logger.info(f"Archivos en directorio: {files}")
                
                # Buscar cualquier PDF generado
                for file in files:
                    if file.endswith(".pdf"):
                        pdf_path = os.path.join(output_dir, file)
                        logger.info(f"PDF encontrado: {pdf_path}")
                        return pdf_path
                
                logger.error("No se encontró ningún PDF generado")
                return None
                
        except Exception as e:
            logger.error(f"Error en conversión: {str(e)}")
            return None
    
    @staticmethod
    def convert_to_pdf(word_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Convierte un documento de Word a PDF con modificación de encabezados
        
        Args:
            word_data: Bytes del documento Word
            filename: Nombre del archivo original
            
        Returns:
            Dictionary with PDF path and filename
        """
        # Crear un directorio temporal para trabajar
        temp_dir = tempfile.mkdtemp()
        logger.info(f"Directorio temporal creado: {temp_dir}")
        
        try:
            # Obtener el nombre base del archivo sin extensión
            base_name = os.path.splitext(filename)[0]
            
            # Definir rutas para los archivos
            word_path = os.path.join(temp_dir, filename)
            pdf_filename = f"{base_name}.pdf"
            
            # Guardar el documento Word en el directorio temporal
            with open(word_path, "wb") as f:
                f.write(word_data)
            
            logger.info(f"Documento Word guardado en: {word_path}")
            
            # Paso 1: Modificar el documento para eliminar encabezados y estandarizar fuentes
            result = WordToPdfService.modify_document_headers(word_path)
            if not result or not result[0]:
                logger.error(f"Error al modificar encabezados en {word_path}")
                raise Exception("Error al procesar el documento")
            
            modified_docx, base_code = result
            
            # Crear un subdirectorio para la salida del PDF
            pdf_output_dir = os.path.join(temp_dir, "pdf_output")
            os.makedirs(pdf_output_dir, exist_ok=True)
            logger.info(f"Directorio para PDF creado: {pdf_output_dir}")
            
            # Paso 2: Convertir a PDF usando LibreOffice
            output_pdf = WordToPdfService.convert_docx_to_pdf(modified_docx, pdf_output_dir)
            
            if not output_pdf:
                logger.error(f"Error al convertir {modified_docx}")
                raise Exception("Error al convertir el documento")
            
            if not os.path.isfile(output_pdf):
                logger.error(f"La ruta de salida no es un archivo válido: {output_pdf}")
                raise Exception("Error: la ruta de salida del PDF no es un archivo válido")
                
            logger.info(f"PDF generado en: {output_pdf}")
            
            # Paso 3: Modificar el PDF para añadir encabezados correctos en cada página
            modified_pdf = WordToPdfService.add_page_headers_to_pdf(output_pdf, base_code)
            
            if not modified_pdf:
                logger.error(f"Error al modificar encabezados en el PDF {output_pdf}")
                raise Exception("Error al modificar encabezados en el PDF")
            
            logger.info(f"Conversión exitosa con encabezados modificados: {modified_pdf}")
            
            # Usar el PDF modificado como resultado final
            output_pdf = modified_pdf
            
            return {
                "pdf_path": output_pdf,
                "filename": pdf_filename
            }
            
        except Exception as e:
            # Limpiar el directorio temporal en caso de error
            shutil.rmtree(temp_dir, ignore_errors=True)
            raise e
